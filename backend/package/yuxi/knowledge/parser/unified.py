"""Unified parser module for markdown conversion."""

from __future__ import annotations

import asyncio
import base64
import os
import re
import tempfile
import threading
import time
from pathlib import Path
from typing import Any

import aiofiles
from docling.backend.msexcel_backend import MsExcelDocumentBackend
from docling.backend.mspowerpoint_backend import MsPowerpointDocumentBackend
from docling.backend.msword_backend import MsWordDocumentBackend
from docling.datamodel.base_models import InputFormat
from docling.datamodel.document import InputDocument
from docling_core.types.doc import DoclingDocument
from markdownify import markdownify as md_convert
from pypdf import PdfReader

from yuxi.knowledge.parser.capabilities import (
    IMAGE_FILE_EXTENSIONS,
    get_ocr_engine_ids,
)
from yuxi.knowledge.parser.zip_utils import process_zip_file as _process_zip_file
from yuxi.knowledge.utils.pdf_utils import validate_pdf_page_tree_loadable
from yuxi.storage.minio import get_minio_client
from yuxi.utils import logger

_OFFICE_BACKENDS = {
    ".docx": (InputFormat.DOCX, MsWordDocumentBackend),
    ".pptx": (InputFormat.PPTX, MsPowerpointDocumentBackend),
    ".xlsx": (InputFormat.XLSX, MsExcelDocumentBackend),
    ".xls": (InputFormat.XLS, MsExcelDocumentBackend),
}
_docling_office_lock = threading.Lock()


def pdfreader(file_path, params=None):
    """读取 PDF 文件并返回 text 文本。"""
    if isinstance(file_path, str):
        file_path = Path(file_path)

    assert file_path.exists(), "File not found"
    assert file_path.suffix.lower() == ".pdf", "File format not supported"

    with file_path.open("rb") as pdf_file:
        reader = PdfReader(pdf_file)
        return "\n\n".join(page.extract_text(extraction_mode="plain").strip() for page in reader.pages)


def parse_pdf(file, params=None):
    """解析 PDF 文件，支持多种 OCR 方式。"""
    from yuxi.knowledge.parser.base import DocumentProcessorException
    from yuxi.knowledge.parser.factory import DocumentProcessorFactory

    opt_ocr, processor_params = _resolve_ocr_engine_params(params)

    if opt_ocr == "disable":
        return pdfreader(file, params=processor_params)

    image_bucket, image_prefix = _resolve_image_storage_params(processor_params)
    processor_params.setdefault("image_bucket", image_bucket)
    processor_params.setdefault("image_prefix", image_prefix)
    processor_kwargs = processor_params.pop("_ocr_processor_kwargs", {})

    try:
        return DocumentProcessorFactory.process_file(opt_ocr, file, processor_params, processor_kwargs)
    except DocumentProcessorException as e:
        logger.error(f"文档处理失败: {e.service_name} - {str(e)}")
        raise
    except Exception as e:  # noqa: BLE001
        logger.error(f"PDF 解析失败: {str(e)}")
        raise DocumentProcessorException(f"PDF解析失败: {str(e)}", opt_ocr, "parsing_failed")


def parse_image(file, params=None):
    """解析图像文件，支持多种 OCR 方式。"""
    from yuxi.knowledge.parser.base import DocumentProcessorException
    from yuxi.knowledge.parser.factory import DocumentProcessorFactory

    opt_ocr, processor_params = _resolve_ocr_engine_params(params)

    if opt_ocr == "disable":
        raise ValueError(
            f"图像文件必须启用OCR才能提取文本内容。请选择OCR方式 ({'/'.join(get_ocr_engine_ids())}) 或移除该文件。"
        )

    image_bucket, image_prefix = _resolve_image_storage_params(processor_params)
    processor_params.setdefault("image_bucket", image_bucket)
    processor_params.setdefault("image_prefix", image_prefix)
    processor_kwargs = processor_params.pop("_ocr_processor_kwargs", {})

    try:
        return DocumentProcessorFactory.process_file(opt_ocr, file, processor_params, processor_kwargs)
    except DocumentProcessorException as e:
        logger.error(f"图像处理失败: {e.service_name} - {str(e)}")
        raise
    except Exception as e:  # noqa: BLE001
        logger.error(f"图像解析失败: {str(e)}")
        raise DocumentProcessorException(f"图像解析失败: {str(e)}", opt_ocr, "parsing_failed")


async def parse_pdf_async(file, params=None):
    return await asyncio.to_thread(parse_pdf, file, params=params)


async def parse_image_async(file, params=None):
    return await asyncio.to_thread(parse_image, file, params=params)


async def parse_resolved_document(source: str, params: dict | None = None) -> str:
    """使用已解析的运行时参数，将本地或 MinIO 文件转换为 Markdown。"""
    from yuxi.knowledge.utils.kb_utils import is_minio_url, parse_minio_url
    from yuxi.storage.minio.client import get_minio_client

    # 1. 如果是 MinIO URL，下载文件到临时路径
    if is_minio_url(source):
        logger.debug(f"Downloading file from MinIO: {source}")

        if "?" in source:
            source_clean = source.split("?")[0]
        else:
            source_clean = source

        original_filename = source_clean.split("/")[-1]

        with tempfile.NamedTemporaryFile(delete=False, suffix=Path(original_filename).suffix) as temp_file:
            temp_path = temp_file.name

        try:
            bucket_name, object_name = parse_minio_url(source)
            minio_client = get_minio_client()
            file_content = await minio_client.adownload_file(bucket_name, object_name)

            async with aiofiles.open(temp_path, "wb") as f:
                await f.write(file_content)

            logger.debug(f"File downloaded to temp path: {temp_path}")
            actual_file_path = temp_path

        except Exception as e:  # noqa: BLE001
            if os.path.exists(temp_path):
                os.unlink(temp_path)
            logger.error(f"Failed to download file from MinIO: {e}")
            raise ValueError(f"无法从MinIO下载文件: {e}")
    else:
        actual_file_path = source

    # 2. 根据文件类型调用不同的解析器
    try:
        file_path_obj = Path(actual_file_path)
        file_ext = file_path_obj.suffix.lower()

        if file_ext == ".pdf":
            validate_pdf_page_tree_loadable(file_path_obj)
            result = await parse_pdf_async(str(file_path_obj), params=params)

        elif file_ext in [".txt", ".md"]:
            async with aiofiles.open(file_path_obj, encoding="utf-8") as f:
                content = await f.read()
            result = content

        elif file_ext == ".docx":
            try:
                result = await asyncio.to_thread(_convert_with_docling, file_path_obj, params=params)
            except Exception as e:  # noqa: BLE001
                logger.warning(f"Docling 解析 DOCX 失败，回退到 python-docx: {file_path_obj.name}, {e}")
                result = await asyncio.to_thread(_convert_docx_with_python_docx, file_path_obj)

        elif file_ext == ".pptx":
            result = await asyncio.to_thread(_convert_with_docling, file_path_obj, params=params)

        elif file_ext in IMAGE_FILE_EXTENSIONS:
            result = await parse_image_async(str(file_path_obj), params=params)

        elif file_ext in [".html", ".htm"]:
            async with aiofiles.open(file_path_obj, encoding="utf-8") as f:
                content = await f.read()
            text = await asyncio.to_thread(md_convert, content, heading_style="ATX")
            result = f"{text}"

        elif file_ext == ".csv":
            result = await asyncio.to_thread(_convert_csv_to_markdown, file_path_obj)

        elif file_ext in [".xls", ".xlsx"]:
            result = await asyncio.to_thread(_convert_with_docling, file_path_obj, params=params)

        elif file_ext == ".json":
            import json

            async with aiofiles.open(file_path_obj, encoding="utf-8") as f:
                content = await f.read()
            data = json.loads(content)
            json_str = json.dumps(data, ensure_ascii=False, indent=2)
            result = f"```json\n{json_str}\n```"

        elif file_ext == ".zip":
            image_bucket, image_prefix = _resolve_image_storage_params(params)
            result = await _process_zip_file(
                str(file_path_obj),
                image_bucket=image_bucket,
                image_prefix=image_prefix,
            )

        else:
            raise ValueError(f"Unsupported file type: {file_ext}")

    except Exception:
        if is_minio_url(source) and os.path.exists(actual_file_path):
            try:
                os.unlink(actual_file_path)
                logger.debug(f"Cleaned up temp file: {actual_file_path}")
            except Exception as cleanup_e:  # noqa: BLE001
                logger.warning(f"Failed to clean up temp file {actual_file_path}: {cleanup_e}")
        raise

    finally:
        if is_minio_url(source) and os.path.exists(actual_file_path):
            try:
                os.unlink(actual_file_path)
                logger.debug(f"Cleaned up temp file: {actual_file_path}")
            except Exception as e:  # noqa: BLE001
                logger.warning(f"Failed to clean up temp file {actual_file_path}: {e}")

    return result


def _convert_office_document(file_path: Path) -> DoclingDocument:
    """用 Docling Slim 的格式 backend 转换一个 Office 文件。"""
    try:
        input_format, backend_type = _OFFICE_BACKENDS[file_path.suffix.lower()]
    except KeyError as exc:
        raise ValueError(f"Docling 不支持该 Office 格式: {file_path.suffix}") from exc

    input_document = InputDocument(file_path, input_format, backend_type)
    backend = getattr(input_document, "_backend", None)
    if backend is None:
        raise RuntimeError(f"Docling 无法读取 Office 文件: {file_path.name}")

    try:
        if not input_document.valid:
            raise RuntimeError(f"Docling 无法读取 Office 文件: {file_path.name}")
        return backend.convert()
    finally:
        backend.unload()


def _resolve_image_storage_params(params: dict | None) -> tuple[str, str]:
    params = params or {}

    image_bucket = params.get("image_bucket") or get_minio_client().KB_BUCKETS["images"]
    image_prefix = params.get("image_prefix")
    if image_prefix:
        normalized_prefix = str(image_prefix).strip("/")
        if normalized_prefix:
            return image_bucket, normalized_prefix

    return image_bucket, "unknown/kb-images"


def _resolve_ocr_engine_params(params: dict | None) -> tuple[str, dict[str, Any]]:
    params = params or {}
    engine = str(params.get("ocr_engine") or "").strip()
    if not engine:
        raise ValueError("OCR 文件缺少已解析的 ocr_engine，请通过 parse_document() 解析")

    processor_params = dict(params)
    processor_params.pop("ocr_engine_config", None)
    return engine, processor_params


def _upload_image_to_minio(image_data: bytes, filename: str, bucket_name: str, object_prefix: str) -> str:
    """上传图片到 MinIO，返回经后端鉴权代理访问的 URL。"""
    from yuxi.knowledge.utils.kb_utils import build_kb_image_proxy_url

    minio_client = get_minio_client()
    minio_client.ensure_bucket_exists(bucket_name)

    normalized_prefix = object_prefix.strip("/") or "unknown/kb-images"
    timestamp = int(time.time() * 1000000)
    object_name = f"{normalized_prefix}/{timestamp}_{Path(filename).name}"

    minio_client.upload_file(
        bucket_name=bucket_name,
        object_name=object_name,
        data=image_data,
    )
    return build_kb_image_proxy_url(object_name)


def _parse_data_uri(data_uri: str) -> tuple[bytes, str]:
    """解析 data URI，返回 (image_data, mime_type)。"""
    header, base64_data = data_uri.split(",", 1)
    mime_type = header.split(":")[1].split(";")[0]
    image_data = base64.b64decode(base64_data)
    return image_data, mime_type


def _convert_with_docling(file_path: Path, params: dict | None = None) -> str:
    """使用 Docling Slim 将 Office 文件转换为 Markdown。"""
    params = params or {}
    image_bucket, image_prefix = _resolve_image_storage_params(params)

    with _docling_office_lock:
        doc = _convert_office_document(file_path)

    if hasattr(doc, "pictures") and doc.pictures:
        replacements: list[str] = []
        for pic in doc.pictures:
            uri = str(pic.image.uri) if hasattr(pic, "image") and hasattr(pic.image, "uri") else ""
            if uri.startswith("data:"):
                filename = "image"
                try:
                    image_data, mime_type = _parse_data_uri(uri)
                    filename = f"image_{int(time.time() * 1000000)}.{mime_type.split('/')[-1]}"
                    url = _upload_image_to_minio(image_data, filename, image_bucket, image_prefix)
                    replacements.append(f"![{filename}]({url})")
                except Exception as e:  # noqa: BLE001
                    logger.error(f"上传图片失败 {filename}: {e}")
                    replacements.append(f"[图片: {filename}]")
            else:
                replacements.append("")

        markdown = doc.export_to_markdown()
        for replacement in replacements:
            markdown = re.sub(r"<!--\s*image\s*-->", replacement, markdown, count=1)
        return markdown

    return doc.export_to_markdown()


def _convert_docx_with_python_docx(file_path: Path) -> str:
    """使用 python-docx 解析 DOCX（Docling 失败时兜底）。"""
    from docx import Document

    document = Document(str(file_path))
    blocks: list[str] = []

    for para in document.paragraphs:
        text = para.text.strip()
        if text:
            blocks.append(text)

    for table in document.tables:
        rows: list[list[str]] = []
        for row in table.rows:
            cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
            if any(cells):
                rows.append(cells)

        if not rows:
            continue

        header = rows[0]
        blocks.append(f"| {' | '.join(header)} |")
        blocks.append(f"| {' | '.join(['---'] * len(header))} |")

        for row in rows[1:]:
            normalized_row = row + [""] * (len(header) - len(row))
            blocks.append(f"| {' | '.join(normalized_row[: len(header)])} |")

        blocks.append("")

    return "\n\n".join(blocks).strip()


def _convert_csv_to_markdown(file_path: Path) -> str:
    import pandas as pd

    dataframe = pd.read_csv(file_path)
    tables: list[str] = []
    for i in range(len(dataframe)):
        row_dataframe = dataframe.iloc[[i]]
        tables.append(row_dataframe.to_markdown(index=False))
    return "\n\n".join(tables)
